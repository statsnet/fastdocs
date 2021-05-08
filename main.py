from docx import Document
from docx.shared import Pt
from docx2pdf import convert
from flask_cors import CORS
import os
from flask import Flask
from flask import request, jsonify
import json
import subprocess


app = Flask(__name__)


CORS(app)

def get_templates():
    files = os.listdir(path="templates")
    templates = list()
    for f in files:
        templates.append(f'templates/{f}')
    return templates


def generate_document(filepath, variables):
    doc = Document(filepath)
    style_special = doc.styles['Normal']
    font = style_special.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    for paragraph in doc.paragraphs:
        for key, value in variables.items():
            if key in paragraph.text:
                text = paragraph.text.replace(key, value)
                style = paragraph.style

                paragraph.text = text
                paragraph.style = style

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in variables.items():
                        if key in paragraph.text:
                            text = paragraph.text.replace(key, value)
                            style = paragraph.style

                            paragraph.text = text
                            paragraph.style = style

    filename = filepath.replace('templates/', '').replace('.docx', '')
    doc.save(f'documents/{filename}.docx')
    subprocess.check_output(['libreoffice', '--convert-to', 'pdf', f'documents/{filename}.pdf'])
    return filename + '.pdf'


def order_documents(request):
    results = []
    for template in request['templates']:
        filename = generate_document(template, request['variables'])
        url = 'https://fastdocs.kz/documents/' + filename
        results.append({
            'filename': filename,
            'url': url
        })
    return results


order = {
    'variables': {
            '{{customer.name}}': 'ТОО Обработка больших данных',
            '{{customer.uin}}': '180640032599',
            '{{customer.document}}': 'Устав',
            '{{customer.company_type}}': 'ТОО',
            '{{customer.responsible}}': 'Иванов Иван Иванович',
            '{{customer.responsible_role}}': 'Директор',
            '{{customer.address}}': 'г. Алматы, ул. Жаркоова 126А',
            '{{customer.phone}}': '+77072125198',
            '{{customer.email}}': 'rocketspace@gmail.com',
            '{{customer.iban}}': 'KZ413294023143214123',
            '{{customer.bank}}': 'АО ДБ "Альфа-Банк"',
            '{{customer.bic}}': 'ALFAKZKA',
            '{{holder.uin}}': '201140034418',
            '{{holder.name}}': 'АО Сбербанк',
            '{{holder.document}}': 'Устав',
            '{{holder.company_type}}': 'АО',
            '{{holder.responsible}}': 'Байдильда Нурлан',
            '{{holder.responsible_role}}': 'Генеральный директор',
            '{{contract.number}}': '12-43А',
            '{{contract.date}}': '12.10.2021',
            '{{contract.trademark_name}}': 'Сберсистемс',
            '{{contract.payment_term}}': 'Вариант 1. Предоплата в размере 100%',
            '{{plan.name}}': 'Тариф "Корпоративный"',
            '{{plan.period}}': '365 дней',
            '{{plan.price}}': '1341201',
            '{{plan.amount}}': '1',
            '{{plan.cost}}': '1341201',
            '{{plan.cost_letters}}': 'сто тридцать тысяч тенге'
    },
    'templates': ['templates/1.Основной ТЗ Стандарт RU (1).docx']
}

# print(get_templates())
# print(order_documents(request=order))
# print(json.dumps(order))


@app.route('/')
def home():
    return jsonify({'status': 'OK'})


@app.route('/api/templates', methods=['GET'])
def list_of_templates():
    return jsonify(get_templates())


@app.route('/api/generate_documents', methods=['POST'])
def generate():
    content = request.get_json()
    data = order_documents(content)
    return jsonify(data)


if __name__ == "__main__":
    app.run(host='0.0.0.0', port=5000, debug=True)

